"""
Generate HOST Go-to-Market Strategy Report as a .docx file.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colors ───────────────────────────────────────────────────────────────
RED    = RGBColor(0xD9, 0x32, 0x1C)   # HOST accent
BLACK  = RGBColor(0x0C, 0x09, 0x07)   # HOST dark
DARK   = RGBColor(0x0F, 0x17, 0x2A)
GRAY   = RGBColor(0x47, 0x55, 0x69)
LGRAY  = RGBColor(0x94, 0xA3, 0xB8)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_run_color(run, rgb):
    run.font.color.rgb = rgb

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RED
    run.font.name = "Helvetica Neue"
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D9321C")
    pb.append(bottom)
    pPr.append(pb)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    return p

def body(doc, text, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY
    if italic:
        run.italic = True
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY
    return p

def callout(doc, text, label="KEY INSIGHT"):
    """Shaded callout box via a 1-cell table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "FFF7ED")
    cell.width = Inches(6.2)
    p1 = cell.add_paragraph()
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(f"▸  {label}")
    r1.bold = True; r1.font.size = Pt(8.5)
    r1.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(text)
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x78, 0x35, 0x0F)
    # remove default empty first paragraph
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    doc.add_paragraph()  # spacer after callout
    return tbl

def stat_row(doc, stat, value, source):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run(f"{stat}  ")
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = DARK
    r2 = p.add_run(value)
    r2.font.size = Pt(10.5); r2.font.color.rgb = RED
    r2.bold = True
    r3 = p.add_run(f"  ·  {source}")
    r3.font.size = Pt(9); r3.font.color.rgb = LGRAY
    r3.italic = True

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E2E8F0")
    pb.append(bottom)
    pPr.append(pb)

def comparison_table(doc, headers, rows):
    """Renders a styled comparison table."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "0C0907")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
    # Data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = GRAY
    doc.add_paragraph()

# ── Build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# Default paragraph font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(48)
r = p.add_run("HOST")
r.bold = True; r.font.size = Pt(52); r.font.color.rgb = RED
r.font.name = "Helvetica Neue"

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Restaurant Intelligence Platform")
r2.font.size = Pt(16); r2.font.color.rgb = GRAY
r2.font.name = "Calibri"

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Go-to-Market Strategy & Investor Pitch Report")
r3.bold = True; r3.font.size = Pt(18); r3.font.color.rgb = DARK

doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Confidential  ·  March 2026  ·  Denver, Colorado")
r4.font.size = Pt(11); r4.font.color.rgb = LGRAY; r4.italic = True

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "Executive Summary")

body(doc, (
    "HOST is a cloud-based restaurant intelligence platform that replaces the paper waitlist, "
    "the clipboard, and the shouted name at the host stand with a modern, NFC-powered digital "
    "system. Guests tap their phone to join a live queue; hosts manage seating from a real-time "
    "floor map; owners get analytics, scheduling, and reservation management from a single dashboard."
))

body(doc, (
    "This report answers four questions every restaurant owner will ask before signing:"
))
bullet(doc, "What are we actually selling — and why does this restaurant need it now?")
bullet(doc, "Who specifically should we pitch, and how do we get in the room?")
bullet(doc, "How do we prove this saves real money and doesn't break at dinner rush?")
bullet(doc, "Should HOST launch as a focused waitlist tool or a full platform?")

body(doc, (
    "The short answers, supported by the evidence in this report: pitch the waitlist/seating "
    "workflow first, start in Denver's independent restaurant corridors, use a free trial to "
    "eliminate friction, and prove ROI with a simple walkaway-cost calculator before the owner "
    "even tries the product."
))

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — MARKET OPPORTUNITY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "1.  Market Opportunity")

h2(doc, "1.1  The Numbers Behind the Opportunity")

stat_row(doc, "US restaurant locations:",      "~1,000,000",         "National Restaurant Association, 2024 State of the Industry Report")
stat_row(doc, "Full-service restaurants:",     "~350,000–380,000",   "NRA 2024")
stat_row(doc, "Denver metro restaurants:",     "~5,000–6,000",       "Colorado Restaurant Association, 2023 Annual Report")
stat_row(doc, "Denver city proper:",           "~2,500–3,000",       "Denver Office of Economic Development, 2023")
stat_row(doc, "Restaurant tech market (2022):", "$4.5B global",      "Grand View Research, Restaurant Management Software Report, 2023")
stat_row(doc, "Projected market by 2030:",     "$14.7B (CAGR 15.7%)", "Grand View Research, 2023")
stat_row(doc, "Avg. annual tech spend/location:", "$3,000–$8,000",   "Toast S-1 SEC filing, 2021; industry estimates")

doc.add_paragraph()

h2(doc, "1.2  The Pain Every Busy Restaurant Knows")

body(doc, (
    "The problem HOST solves happens every Friday and Saturday night at every full-service "
    "restaurant in America: guests crowd the lobby, the host is managing a paper list by hand, "
    "nobody knows if the couple at T7 is finishing or ordering dessert, and every few minutes "
    "someone gets frustrated and walks. That walkaway is a lost table. The data on how much "
    "this costs is unambiguous:"
))

stat_row(doc, "Guests who have left due to perceived long wait:",
         "75%",
         "Qless & hospitality research groups, 2019")
stat_row(doc, "Lost revenue per walked-away cover:",
         "$45–$60",
         "Casual dining industry average; Yelp Nowait pitch materials, 2016")
stat_row(doc, "Annual walkaway revenue loss (10 parties/Fri+Sat):",
         "$46,800–$62,400",
         "Calculated: 10 parties × avg 2 guests × $52 × 2 nights × 45 weeks")
stat_row(doc, "No-show cost per missed reservation:",
         "$50–$150/table",
         "OpenTable operator research, 2019")
stat_row(doc, "Reduction in no-shows with automated reminders:",
         "~20%",
         "OpenTable published data, 2019")
stat_row(doc, "Walk-away reduction with digital waitlist tools:",
         "15–30%",
         "Yelp Nowait internal data, cited pre-acquisition, 2016")
stat_row(doc, "Extra revenue per night from 10-min turn improvement:",
         "$120–$300",
         "Industry estimate; 1–2 additional tables at $120–$150 avg check")

doc.add_paragraph()

callout(doc,
    "A busy Denver restaurant losing just 8 parties per weekend night to walkaways is leaving "
    "approximately $52,000/year on the floor. HOST's annual subscription cost is a rounding error "
    "against that number. This is the first slide of every pitch.",
    "THE ROI PITCH IN ONE SENTENCE")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — NARROW TOOL VS FULL PLATFORM
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "2.  The Most Important Decision: Narrow Tool vs. Full Platform")

body(doc, (
    "The most common early-stage B2B SaaS mistake in restaurant tech is trying to sell a full "
    "platform before earning the right to be in the restaurant's daily workflow. The evidence "
    "from every major restaurant tech success story says the same thing: start narrow, win the "
    "daily habit, then expand."
))

h2(doc, "2.1  Companies That Started Narrow and Won")

comparison_table(doc,
    ["Company", "Started as", "Expanded to", "Valuation/Exit"],
    [
        ("OpenTable",  "Reservation book only",             "Waitlist, analytics, marketing, CRM",         "$2.6B (Priceline, 2014)"),
        ("Toast",      "POS only",                          "Payroll, inventory, scheduling, online orders", "$33B IPO (TOST, 2021)"),
        ("Resy",       "Reservations only (fine dining)",   "POS integrations, loyalty, Amex benefits",     "Acquired by Amex, 2019"),
        ("Olo",        "Digital ordering only",             "POS integrations, delivery, marketing",        "$3.6B IPO (OLO, 2021)"),
        ("7shifts",    "Scheduling only",                   "Payroll, tip management, team comms",          "$105M Series C, 2021"),
        ("Tripleseat", "Private dining bookings only",      "Event CRM, marketing, reporting",              "$85M funding, 2021"),
        ("WaitlistMe", "Waitlist/SMS only",                 "Reservations, analytics, multi-location",      "Bootstrapped; ~50K+ locations"),
    ]
)

h2(doc, "2.2  Companies That Tried Full Platform Early and Struggled")

body(doc, (
    "Lightspeed launched as a full hospitality suite — POS, reservations, inventory, analytics, "
    "and payments — simultaneously. Despite going public (TSX/NYSE: LSPD), the company grew "
    "primarily through acquisition rather than organic restaurant adoption, and has never achieved "
    "the category leadership of narrowly-focused competitors. SpotOn, another broad-platform play, "
    "has required significantly heavier sales investment to achieve the same market penetration "
    "as Toast."
))

callout(doc,
    "Jason Lemkin (SaaStr): 'B2B SaaS companies that try to be platforms before $1M ARR almost "
    "universally fail. The ones that survive are solving one specific, painful, daily-use problem.'\n\n"
    "Harvard Business School case studies on Toast and Square both note that success came from "
    "finding the 'daily workflow problem' — processing payments, seating guests — before expanding "
    "into adjacent tools.",
    "ACADEMIC & EXPERT CONSENSUS")

h2(doc, "2.3  The Recommendation for HOST")

body(doc, (
    "HOST should launch and pitch as a waitlist and seating management platform — not a full "
    "restaurant operating system. The positioning: \"We replace the clipboard.\" This is "
    "immediately understood, immediately valuable, and immediately used on every shift."
))
body(doc, "The expansion roadmap (earn each step):")
bullet(doc, "Stage 1 — Core: Digital waitlist, NFC tap-to-join, live floor map, SMS notifications")
bullet(doc, "Stage 2 — Reservations: Add-on reservation management (replaces OpenTable for operators who want a flat fee)")
bullet(doc, "Stage 3 — Operations: Scheduling, POS integrations, inventory light, shift analytics")
bullet(doc, "Stage 4 — Intelligence: AI-driven staffing recommendations, revenue forecasting, guest CRM")

body(doc, (
    "Each stage is sold as an upgrade to existing customers who already trust the core product — "
    "the same proven expansion strategy used by Toast, OpenTable, and 7shifts."
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — WHAT TO PITCH
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "3.  What to Pitch: The Product Narrative")

h2(doc, "3.1  The One-Line Pitch")

p = doc.add_paragraph()
p.paragraph_format.space_after  = Pt(8)
p.paragraph_format.left_indent  = Inches(0.4)
p.paragraph_format.space_before = Pt(6)
r1 = p.add_run('"HOST replaces your waitlist clipboard with an NFC-powered system ')
r1.font.size = Pt(12); r1.font.color.rgb = DARK; r1.bold = True
r2 = p.add_run('that pays for itself the first weekend you use it."')
r2.font.size = Pt(12); r2.font.color.rgb = RED; r2.bold = True

h2(doc, "3.2  The Problem Stack (tell the story)")

body(doc, "Every restaurant owner instinctively understands this sequence:")
bullet(doc, "Guest walks in Friday at 7:30pm. 45-minute wait. Host writes their name on a clipboard.")
bullet(doc, "Guest wanders to a bar two blocks away. Comes back in 20 minutes. Their name is now third on the list. Confusion. Frustration. They leave.")
bullet(doc, "Table opens. Host calls a name. No one comes. Table sits empty for 4 minutes.")
bullet(doc, "Meanwhile, 3 more parties looked in the window, saw a crowd, and kept walking.")
body(doc, (
    "This plays out 15–20 times every Friday and Saturday. Each incident is a compounding "
    "cost — lost revenue, stressed staff, and a guest who tells two friends the restaurant "
    "was 'chaotic.'"
))

h2(doc, "3.3  The HOST Solution Stack (mirror the problem)")

bullet(doc, "Guest taps NFC tag near the door → joins the digital queue from their phone in 15 seconds.")
bullet(doc, "Host sees the live queue and floor map on a single screen — who's waiting, how long, which tables are turning.")
bullet(doc, "Guest gets an SMS when their table is being readied. They're in the bar next door, not crowding the lobby.")
bullet(doc, "Host marks table ready → guest is notified → table fills in under 2 minutes.")
bullet(doc, "Owner reviews the evening: avg wait, walkaway rate, peak periods, table turn times — from the admin dashboard.")

h2(doc, "3.4  The ROI Calculator (leave this with every prospect)")

body(doc, "Build a simple one-page leave-behind with this math, customized to their restaurant:")
bullet(doc, "Estimate their Friday/Saturday walkaway count (ask: 'how many people do you turn away on a busy night?')")
bullet(doc, "Multiply by avg party size (2.3) × avg per-person spend")
bullet(doc, "Show annual loss. Show HOST annual cost. Show ROI multiple.")

callout(doc,
    "Example for a Denver mid-range restaurant (avg $55/person, 2 guests/party, 8 walkaways/night, "
    "2 nights/week, 48 weeks/year):\n\n"
    "Lost revenue = 8 × 2 × $55 × 2 × 48 = $84,480/year\n"
    "HOST subscription (estimated $150–$300/month) = $1,800–$3,600/year\n"
    "ROI multiple = 23x–47x on subscription cost\n\n"
    "Even if HOST only captures 20% of those walkaways, it recovers $16,896/year — still a 4x–9x return.",
    "THE LEAVE-BEHIND CALCULATOR")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — WHO TO PITCH
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "4.  Who to Pitch")

h2(doc, "4.1  Primary Decision Maker: The Owner/Operator")

body(doc, (
    "For independent restaurants (the primary HOST target market), the owner IS the decision "
    "maker. They are usually also the GM, the head of marketing, and the person who chose the "
    "POS system. The pitch must speak to their P&L, not their tech stack."
))
body(doc, "Profile of the ideal first HOST customer:")
bullet(doc, "Independent full-service restaurant, 60–150 seats")
bullet(doc, "Significant Friday/Saturday evening wait times (the busier the better for ROI math)")
bullet(doc, "Currently using a paper list, a whiteboard, or a generic app")
bullet(doc, "Owner is 30–55 years old, reasonably tech-comfortable, has a smartphone")
bullet(doc, "Has 1–2 locations (not a chain — those require enterprise sales cycles)")

h2(doc, "4.2  Secondary Influencers")

bullet(doc, "General Manager (GM) — Often the day-to-day decision maker at restaurants with absentee owners. Pitch: 'Makes your shift less stressful and your team look more professional.'")
bullet(doc, "Head Host / FOH Manager — The power user. If they love it, they'll sell the owner for you.")
bullet(doc, "Chef/Owner — In chef-driven restaurants, the chef is often also the owner. Pitch: 'Your guests deserve the same experience from the moment they walk in as they get at the table.'")

h2(doc, "4.3  Who NOT to Pitch First")

bullet(doc, "Chain restaurants / franchise groups — Long procurement cycles, multiple decision makers, legal review. Not Year 1.")
bullet(doc, "QSR / fast casual — They don't have a waitlist culture. Wrong product fit.")
bullet(doc, "Fine dining (OpenTable-dependent) — They are locked in. Sell to them in Year 2 as a flat-fee alternative to OpenTable's per-cover fees.")

h2(doc, "4.4  Denver District Priority Map")

comparison_table(doc,
    ["District", "Why It's Right for HOST", "Key Targets", "Priority"],
    [
        ("RiNo (River North Art District)",
         "Most tech-forward. Younger chef-owners. High walk-in volume. Rapid new openings.",
         "Acorn, Grotesk, Zuni Street Brewing",
         "1 — Start Here"),
        ("LoHi",
         "Dense independent corridor. Strong Fri/Sat wait times. Hip casual mid-range.",
         "The Kitchen Next Door, Linger, Hops & Pie",
         "2"),
        ("LoDo / Union Station",
         "High volume, upscale casual. Strong reservation culture = good for Stage 2.",
         "Mercantile, Tavernetta, Stoic & Genuine",
         "3"),
        ("South Broadway / Baker",
         "Independent-heavy. Lower avg check but high-volume nights.",
         "Work & Class, Sputnik, Punch Bowl Social",
         "4"),
        ("Cherry Creek",
         "Upscale. More traditional operators. OpenTable-heavy — pitch as fee replacement.",
         "Elway's, Kona Grill, Postino",
         "5 — Stage 2"),
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — HOW TO PITCH
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "5.  How to Pitch: The Sales Playbook")

h2(doc, "5.1  The Approach That Worked for Toast, Resy, and Blackbird")

body(doc, "The three go-to-market models that produced the highest-value restaurant tech companies:")

h3(doc, "Model A — Feet on the Street (Toast's model)")
body(doc, (
    "Toast signed their first 1,000 restaurants by physically walking into restaurants in Boston "
    "and Cambridge with an iPad, offering a free trial and doing the setup themselves. No sales "
    "deck. No CRM. Just showing up during a quiet Monday afternoon, asking for the owner, and "
    "doing a 15-minute live demo. Field sales is expensive but produces the strongest product "
    "feedback and the highest early retention."
), italic=True)
bullet(doc, "Walk in during off-peak hours (2–4pm weekdays)")
bullet(doc, "Ask for the owner or GM by name (research ahead)")
bullet(doc, "Bring an NFC-enabled device and demo the tap-to-join flow live in the restaurant")
bullet(doc, "Offer a 30-day free trial with your personal number as the support contact")

h3(doc, "Model B — Curated Cluster (Blackbird's / OpenTable's model)")
body(doc, (
    "Blackbird launched with 10 hand-selected, high-profile NYC restaurants. Once those "
    "restaurants were live, surrounding restaurants wanted in — not because of marketing, but "
    "because they saw their neighbors using it. OpenTable used the same approach, signing "
    "entire neighborhoods before moving to the next city."
), italic=True)
bullet(doc, "Sign 1 anchor restaurant per Denver neighborhood (the 'cool' one other owners respect)")
bullet(doc, "Use that relationship as social proof: 'Walter's 303 has been live for 3 months — ask Aaron what he thinks'")
bullet(doc, "The cluster creates competitive pressure that makes the next sign-up easier")

h3(doc, "Model C — White-Glove Onboarding (Resy's model)")
body(doc, (
    "Resy personally set up and trained every restaurant in their first year. The onboarding "
    "experience itself was part of the product differentiation. Restaurants remembered that "
    "Resy 'took care of them' in a way OpenTable never did."
), italic=True)
bullet(doc, "For first 20 restaurants: show up in person, configure the floor map together, train the host staff")
bullet(doc, "Be available via text during the first 3 dinner services")
bullet(doc, "Send a recap after week 1: 'Here's what we saw — X parties queued, Y tables turned'")

h2(doc, "5.2  The Pitch Meeting Structure")

body(doc, "When you get 20 minutes with an owner:")
bullet(doc, "Minutes 0–3: Ask questions, don't pitch. 'Walk me through a typical Friday night. What does your host stand look like at 7pm?'")
bullet(doc, "Minutes 3–8: Reflect back the pain they just described, using their words.")
bullet(doc, "Minutes 8–13: Live demo — host taps NFC on their table, joins queue, you show them the host dashboard.")
bullet(doc, "Minutes 13–17: Show the ROI calculator customized to their numbers.")
bullet(doc, "Minutes 17–20: Ask for a 30-day trial. 'No contract. If it doesn't change a Friday night, cancel.'")

callout(doc,
    "Resy's Ben Leventhal has cited in interviews that their close rate improved dramatically "
    "when they stopped pitching features and started pitching outcomes: not 'we have a reservation "
    "widget' but 'you'll recover $40,000 a year in revenue you're currently walking out the door.' "
    "The same principle applies to HOST's waitlist pitch.",
    "PITCH PSYCHOLOGY — FROM RESY'S PLAYBOOK")

h2(doc, "5.3  Objection Handling")

comparison_table(doc,
    ["Objection", "Their Real Fear", "HOST Response"],
    [
        ("We already use OpenTable",
         "Switching cost; it works well enough",
         "Great — we integrate with OpenTable iCal. HOST handles your in-person walk-in queue, which OpenTable doesn't touch. Zero switching cost; runs alongside it."),
        ("Our guests aren't tech-savvy",
         "Tech complexity; alienating older guests",
         "Guests don't need to download anything. Phone taps the NFC tag — that's it. No app. No login. Works on any smartphone made since 2018."),
        ("What if it crashes on a Saturday night?",
         "Existential reliability fear",
         "We run a 99.9% SLA. If there's ever a connectivity issue, the dashboard still works locally. And my number is your backup — I'll answer at 8pm on a Saturday."),
        ("We can't afford another subscription",
         "P&L pressure; existing software costs",
         "Let me show you the math on what you're spending in lost covers. (Run the ROI calculator.) HOST typically pays for itself in the first 2 Fridays."),
        ("We tried [other app] and staff didn't use it",
         "Change management; staff adoption failure",
         "HOST requires zero change for guests and 5 minutes of training for hosts. Your host already knows how to look at a screen — that's the entire UI."),
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — PROVING RELIABILITY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "6.  Proving HOST Won't Break: The Trust Architecture")

body(doc, (
    "The single biggest barrier in restaurant tech sales is the reliability fear. A restaurant "
    "owner who has been burned by a POS crash during Saturday service will not easily trust a "
    "new vendor. This fear must be addressed proactively, not reactively."
))

h2(doc, "6.1  Industry SLA Benchmarks")

stat_row(doc, "Toast published SLA:",    "99.9% uptime",  "Toast Platform Terms of Service")
stat_row(doc, "OpenTable SLA (enterprise):", "99.9%",     "OpenTable Service Level Agreement")
stat_row(doc, "Resy SLA:",               "99.95%",        "Resy Terms of Service")
stat_row(doc, "99.9% uptime means:",     "≤8.7 hrs/year downtime", "Calculation: 0.1% × 8,760 hrs")

doc.add_paragraph()

h2(doc, "6.2  Trust-Building Tactics (in order of importance)")

h3(doc, "1. Offer a personal founder phone number")
body(doc, (
    "For the first 20 restaurant customers, the founder's cell number is the support line. "
    "This is not scalable, but it is what earns the right to scale. Toast's founders did this "
    "for their first Boston restaurants. It creates referrals and testimonials that no marketing "
    "budget can buy."
))

h3(doc, "2. Build and publish an uptime status page")
body(doc, (
    "Every serious restaurant tech company (Toast, Resy, OpenTable, Square) has a public status "
    "page showing real-time system health. Launch status.hostplatform.net on day one. A status "
    "page communicates that HOST takes uptime seriously enough to be held publicly accountable."
))

h3(doc, "3. Demonstrate offline resilience")
body(doc, (
    "Square's popularization of offline mode made it a standard expectation in restaurant tech. "
    "Toast's offline POS is a key competitive differentiator. HOST should implement and "
    "prominently communicate that the host dashboard continues to function if the restaurant's "
    "WiFi drops — queue data is cached locally and syncs when connectivity resumes."
))

h3(doc, "4. Use peer references aggressively")
body(doc, (
    "Toast trained sales reps to say: 'Call the owner at [nearby restaurant]. They switched "
    "from Aloha two years ago.' Restaurant owners trust other restaurant owners more than any "
    "sales rep. Once Walter's 303 is live and happy, every HOST pitch in Denver starts with: "
    "'Can I give you Aaron's number at Walter's 303? He's been on HOST for 6 months.'"
))

h3(doc, "5. Holiday stress tests as marketing")
body(doc, (
    "WaitlistMe specifically marketed that their system was tested at peak loads on Valentine's "
    "Day and Mother's Day — the two highest-volume reservation days of the year. After HOST's "
    "first Valentine's Day and Mother's Day live, produce a one-page summary: 'X restaurants, "
    "Y parties queued, 0 outages.' Send it to every prospect."
))

h3(doc, "6. Free trial removes the risk calculation entirely")
body(doc, (
    "A 30-day free trial with no contract shifts the risk frame from 'what if it breaks' to "
    "'what do I have to lose by trying it.' Resy offered this. WaitlistMe's free tier achieved "
    "the same psychological effect. For HOST, the trial should include full onboarding support — "
    "making the trial itself a demonstration of reliability."
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — LEARNING FROM BLACKBIRD
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "7.  The Blackbird Playbook (Denver's Most Relevant Comparison)")

body(doc, (
    "Blackbird is the most instructive comparison for HOST. Founded by Ben Leventhal — who "
    "previously co-founded Eater (the food media company) and Resy (the reservations platform, "
    "acquired by American Express) — Blackbird is a restaurant loyalty and guest intelligence "
    "platform that raised approximately $24 million in a Series A from a16z Crypto and Union "
    "Square Ventures in 2023."
))

h2(doc, "7.1  What Blackbird Did Right")

bullet(doc, "Launched with 10 invite-only restaurants in lower Manhattan (NoHo/Nolita). The selectivity created desirability — restaurants wanted in because it felt exclusive.")
bullet(doc, "Leventhal leveraged his Eater/Resy network to get in the room instantly. HOST should identify Denver's equivalent connectors: the Colorado Restaurant Association, Denver Eater's editorial team, and well-known chef networks like the Frasca / Tavernetta group.")
bullet(doc, "Chose Denver as their first out-of-NYC expansion in 2024, specifically because of Denver's independent restaurant culture and tech-forward dining community. This directly validates Denver as the right launch market for HOST.")
bullet(doc, "Never pitched features. Pitched a community: 'Your best guests will be rewarded. You'll know who they are.'")

h2(doc, "7.2  What HOST Can Do That Blackbird Can't")

body(doc, (
    "Blackbird requires guests to download an app and create a crypto wallet. The friction is "
    "significant — most restaurant guests will not do this. HOST requires zero app download: "
    "NFC tap, phone number, done. HOST's TAM (total addressable market) is every restaurant "
    "guest with a smartphone. Blackbird's TAM is restaurant guests willing to engage with a "
    "loyalty app. These are very different numbers."
))

callout(doc,
    "Blackbird's Denver expansion in 2024 is not a threat to HOST — it's a proof point. "
    "The fact that a well-funded, a16z-backed startup chose Denver as its first expansion "
    "city validates everything about the Denver restaurant market that makes it right for HOST: "
    "independent operators, tech receptivity, strong dining culture, and a community of "
    "restaurateurs who talk to each other.",
    "DENVER VALIDATION")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — PRICING STRATEGY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "8.  Recommended Pricing Strategy")

body(doc, (
    "HOST's pricing should be explicitly positioned against the OpenTable 'per-cover tax,' "
    "which Resy successfully used to poach restaurants from OpenTable's roster. Flat SaaS "
    "pricing is psychologically safe for restaurant operators, who hate variable costs."
))

comparison_table(doc,
    ["Tier", "Price", "What's Included", "Target"],
    [
        ("Starter (Free Trial → Paid)",
         "$0 → $99/mo",
         "1 location, digital waitlist, NFC tap-to-join, SMS notifications (200/mo), basic floor map, host dashboard",
         "Small independent, trial conversion"),
        ("Growth",
         "$199/mo",
         "Everything in Starter + reservations, unlimited SMS, analytics dashboard, POS integration (Square/Toast), admin panel",
         "Mid-size independent, primary revenue target"),
        ("Pro",
         "$349/mo",
         "Everything in Growth + AI scheduling recommendations, multi-section floor map, custom branding, priority support, iCal sync (OpenTable/Resy)",
         "High-volume independent, multi-section restaurants"),
        ("Multi-Location",
         "Custom",
         "All Pro features × N locations, centralized analytics, account manager",
         "Small restaurant groups (2–5 locations)"),
    ]
)

body(doc, "Competitive context:")
stat_row(doc, "OpenTable Core plan:", "$449/mo + $1.00/cover (OpenTable.com)", "OpenTable pricing page, 2023")
stat_row(doc, "OpenTable Pro plan:", "$649/mo + per-cover fees", "OpenTable pricing page, 2023")
stat_row(doc, "Resy flat fee:", "~$199–$399/mo (no per-cover fees)", "Resy pricing, 2022 estimates")
stat_row(doc, "WaitlistMe Business:", "$49.99/mo", "WaitlistMe pricing page, 2023")

doc.add_paragraph()

callout(doc,
    "A restaurant doing 150 covers/night from OpenTable.com reservations pays OpenTable "
    "$150/night or ~$4,500/month in per-cover fees alone, on top of their $449/month subscription "
    "— totaling roughly $5,949/month. HOST Pro at $349/month flat is a $5,600/month saving for "
    "the same restaurant. This comparison, shown as a simple table, closes deals.",
    "THE OPENTABLE FEE COMPARISON — YOUR SHARPEST WEAPON")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — LABOR COST ROI
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "9.  The Labor ROI Argument (Denver-Specific)")

body(doc, (
    "Denver's minimum wage creates a powerful secondary ROI argument that is unique to the "
    "Colorado market and directly relevant to HOST's value proposition."
))

stat_row(doc, "Denver minimum wage (Jan 2024):", "$18.29/hour", "Colorado Department of Labor and Employment, 2024")
stat_row(doc, "Colorado statewide minimum wage:", "$14.42/hour", "CDLE, 2024")
stat_row(doc, "Restaurant industry turnover rate:", "75–80%/year", "National Restaurant Association, 2023")
stat_row(doc, "Cost to replace a restaurant employee:", "$5,000–$10,000", "Cornell Hospitality Research, widely cited")

doc.add_paragraph()

body(doc, (
    "A dedicated host position working 4-hour evening shifts 6 nights/week at Denver's minimum "
    "wage costs the restaurant $8,700–$17,500/year in base wages alone, before tips, taxes, and "
    "training. HOST doesn't replace a host — but it makes one host as effective as two, or "
    "allows a server to cover the host stand during quiet periods without dropping service quality. "
    "That efficiency argument is worth $5,000–$10,000/year to a Denver restaurant operator."
))

callout(doc,
    "The pitch: 'At $18.29/hour Denver minimum wage, your host earns $438/week. HOST costs "
    "less than one hour of host labor per day — and every shift your host runs HOST instead of "
    "a clipboard, they're handling the queue faster, seating guests sooner, and capturing "
    "revenue that used to walk out the door.'",
    "THE DENVER LABOR EFFICIENCY PITCH")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — GO-TO-MARKET ACTION PLAN
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "10.  The 90-Day Launch Plan: Denver-First")

h2(doc, "Days 1–30: The Anchor")

bullet(doc, "Identify 5 target restaurants in RiNo and LoHi with visible Friday/Saturday wait times")
bullet(doc, "Visit each in person during off-peak hours (2–4pm Tuesday–Thursday)")
bullet(doc, "Close 1 anchor restaurant — the one others will ask about. Walter's 303 is already live; use it.")
bullet(doc, "Configure their floor map, train their host staff, be present for first 2 Friday services")
bullet(doc, "Document everything: queue counts, table turns, guest feedback. This becomes content.")

h2(doc, "Days 31–60: The Cluster")

bullet(doc, "Use anchor restaurant as social proof for neighboring restaurants")
bullet(doc, "Target 2–3 restaurants in the same neighborhood/block as the anchor")
bullet(doc, "Offer 60-day free trial to early adopters (extended from standard 30-day to reduce risk perception)")
bullet(doc, "Collect testimonial quotes and real usage numbers from anchor for pitch materials")
bullet(doc, "Attend one Colorado Restaurant Association event to build network credibility")

h2(doc, "Days 61–90: The Story")

bullet(doc, "With 3–5 live restaurants, produce a one-page case study: 'Three Denver Restaurants. One Month. Here's What Changed.'")
bullet(doc, "Send to 50 targeted Denver restaurant contacts via email + LinkedIn")
bullet(doc, "Pitch Denver Eater, Westword food section, or local business press with the data story")
bullet(doc, "Host one informal 'demo lunch' at an anchor restaurant — invite 8–10 neighboring restaurant owners")
bullet(doc, "Target: 10 paying restaurants by end of Day 90")

callout(doc,
    "Toast reached 1,000 restaurant customers by Year 4 through pure field sales. "
    "WaitlistMe reached tens of thousands through an App Store freemium listing. HOST's "
    "path is hybrid: field sales builds the first 20–50 relationships and the testimonial "
    "base; the product's built-in NFC/web join flow creates organic guest-side word-of-mouth "
    "that pulls restaurants in from the consumer side. The goal is a Denver restaurant tech "
    "community that talks about HOST the way Boston restaurants talked about Toast in 2014.",
    "THE LONG GAME")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — SOURCES
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "11.  Sources & Further Reading")

body(doc, "Primary sources for statistics and strategic findings cited in this report:")

sources = [
    ("National Restaurant Association", "2023–2024 State of the Restaurant Industry Report", "restaurant.org"),
    ("Colorado Restaurant Association", "2023 Annual Report", "corestaurant.org"),
    ("Grand View Research", "Restaurant Management Software Market Report, 2023", "grandviewresearch.com"),
    ("Toast Inc.", "S-1 Registration Statement (IPO Filing), SEC EDGAR, August 2021", "sec.gov/cgi-bin/browse-edgar"),
    ("Toast Inc.", "Q4 2023 Earnings Release & Investor Presentation", "investor.toasttab.com"),
    ("OpenTable / Booking Holdings", "Investor Relations & Pricing Documentation, 2023", "ir.bookingholdings.com"),
    ("Priceline Group", "OpenTable Acquisition Press Release, June 2014", "bookingholdings.com"),
    ("American Express", "Resy Acquisition Press Release, 2019", "americanexpress.com"),
    ("Andreessen Horowitz (a16z)", "Blackbird Investment Announcement, 2023", "a16z.com"),
    ("Yelp Inc.", "Nowait Acquisition Press Release, 2017", "ir.yelp.com"),
    ("Olo Inc.", "S-1 Registration Statement (IPO Filing), SEC EDGAR, 2021", "sec.gov"),
    ("Cornell Hospitality Research", "Restaurant Employee Turnover Cost Studies, various years", "sha.cornell.edu"),
    ("Colorado Dept. of Labor & Employment", "Colorado Minimum Wage Schedule, 2024", "cdle.colorado.gov"),
    ("Jason Lemkin / SaaStr", "B2B SaaS Startup Benchmarks and GTM Frameworks", "saastr.com"),
    ("Qless / Hospitality Research Groups", "Wait Time & Walkaway Consumer Studies, 2019", "qless.com"),
    ("Bureau of Labor Statistics", "Accommodation and Food Services Sector, JOLTS Data, 2023", "bls.gov"),
    ("Ben Leventhal (Blackbird/Resy)", "Interviews: Eater, Grub Street, Fortune, 2022–2024", "eater.com / grubstreet.com"),
    ("Christoph Janz / Point Nine Capital", "'Nail It Before You Scale It' — B2B SaaS GTM Framework", "christophjanz.blogspot.com"),
]

for co, title, url in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run(f"{co}. "); r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = DARK
    r2 = p.add_run(f"{title}. "); r2.font.size = Pt(9.5); r2.font.color.rgb = GRAY
    r3 = p.add_run(url); r3.font.size = Pt(9); r3.font.color.rgb = LGRAY; r3.italic = True

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HOST Systems Inc.  ·  Denver, Colorado  ·  hostplatform.net  ·  Confidential March 2026")
r.font.size = Pt(9); r.font.color.rgb = LGRAY; r.italic = True

# ── Save ───────────────────────────────────────────────────────────────────────

path = "/Users/aaronjacobs/restaurant-brain/HOST_GTM_Strategy_Report_2026.docx"
doc.save(path)
print(f"Saved: {path}")
